"""Generate the BADM 576 Week 7 ML Process backup worksheet as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Title
title = doc.add_heading('BADM 576 — Week 7: ML Process In-Class Exercise', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Data Science and Analytics')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Instructions
doc.add_heading('Instructions', level=2)
doc.add_paragraph(
    'In this exercise, you will apply the 7-step ML process to a real business scenario. '
    'Read the case study below, then answer each question thoughtfully. '
    'Explain your reasoning — not just what the answer is, but WHY.'
)

doc.add_paragraph()

# Scenario
doc.add_heading('Business Scenario: HealthPredict', level=2)
doc.add_paragraph(
    'HealthPredict is a health insurance company that wants to predict annual medical costs '
    'for new policyholders at enrollment time. They have 5 years of claims data including '
    'policyholder age, BMI, smoking status, number of dependents, region, and the total '
    'annual medical charges billed.'
)
doc.add_paragraph(
    'Accurate predictions help them set premiums fairly — overpricing loses customers to '
    'competitors, underpricing means the company loses money on high-cost policyholders.'
)

doc.add_paragraph()
doc.add_paragraph('').add_run('_' * 80)

# Questions
questions = [
    {
        'title': 'Step 1: Task Definition (T)',
        'points': 15,
        'questions': [
            'Is this a regression or classification problem? Explain why.',
            'What exactly is the machine learning to predict? Be precise about the target variable, its unit, and the timing of the prediction.',
        ],
    },
    {
        'title': 'Step 2: Data / Experience (E)',
        'points': 20,
        'questions': [
            'What is the outcome variable (Y)?',
            'What predictor variables (X) are available at enrollment time? List them.',
            'Identify at least TWO data quality concerns (think: completeness, accuracy, recency, coverage).',
            'Are there any data leakage risks? What variables might accidentally contain information about the future?',
        ],
    },
    {
        'title': 'Step 3: Model Building',
        'points': 15,
        'questions': [
            'Write the linear regression equation using the predictors you identified above. Use the form: Y = β₀ + β₁X₁ + β₂X₂ + ...',
            'For EACH beta coefficient, state whether you expect it to be positive (+) or negative (−), and explain why.',
        ],
    },
    {
        'title': 'Step 4: Loss Function (P)',
        'points': 15,
        'questions': [
            'What does SSE (Sum of Squared Errors) measure? What about RMSE?',
            'For HealthPredict specifically, which prediction error is MORE costly to the business: overestimating medical costs or underestimating them? Explain your reasoning.',
            'Should the loss function be symmetric or asymmetric for this problem? Why?',
        ],
    },
    {
        'title': 'Step 5: Gradient Descent',
        'points': 10,
        'questions': [
            'Explain conceptually how gradient descent finds the optimal beta values. (No math needed — describe the process in plain English.)',
            'What role does the learning rate play?',
            'Why can\'t we just try every possible combination of beta values?',
        ],
    },
    {
        'title': 'Step 6: Inference / Prediction',
        'points': 10,
        'questions': [
            'Suppose the model has learned the following beta values:\n'
            '    • Intercept (β₀) = 5,000\n'
            '    • Age (β₁) = 250 per year\n'
            '    • BMI (β₂) = 300 per unit\n'
            '    • Smoker (β₃) = 18,000 (1 = yes, 0 = no)\n'
            '    • Dependents (β₄) = 500 per dependent\n\n'
            'A new policyholder enrolls: age = 35, BMI = 28, smoker = yes, dependents = 2.\n\n'
            'Show your work and compute the predicted annual medical cost.',
            'Pick TWO of the beta coefficients above and interpret them in business terms. What does each one mean for HealthPredict?',
        ],
    },
    {
        'title': 'Step 7: Drift Analysis',
        'points': 15,
        'questions': [
            'Give TWO specific, realistic examples of real-world changes that could cause this model to become inaccurate over time.',
            'For each example, state whether it is DATA DRIFT (the input distribution changes) or CONCEPT DRIFT (the relationship between inputs and outcome changes). Explain why.',
        ],
    },
]

for q in questions:
    doc.add_paragraph()
    heading = doc.add_heading(f"{q['title']}  ({q['points']} points)", level=2)

    for i, question in enumerate(q['questions'], 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {question}')
        run.font.size = Pt(11)

        # Answer space
        for _ in range(5):
            doc.add_paragraph()

        # Separator
        sep = doc.add_paragraph()
        sep.add_run('_' * 80).font.color.rgb = RGBColor(200, 200, 200)

# Total
doc.add_paragraph()
total = doc.add_heading('Total: 100 points', level=2)
total.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Name: _______________________     Date: _______________')
run.font.size = Pt(12)

# Save
output_path = r'C:\Users\ashishk\Dropbox\My PC (BUS-P10E67720)\Documents\Development\LLM_Role_Player\BADM576_Week7_ML_Process_Exercise.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
